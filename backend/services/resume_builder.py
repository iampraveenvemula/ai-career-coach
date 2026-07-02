"""
LLM-powered resume builder service.

Uses Ollama to restructure + keyword-align the candidate's resume for a target role,
then exports it as a formatted DOCX using python-docx.
"""
import re
import os
import io
import httpx
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
from services.ats_scorer import calculate_industry_score

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.2")

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

STRUCTURE_PROMPT = """You are a Senior Technical Recruiter at a top-tier technology company (like OpenAI, Google, or Stripe). You only select resumes that showcase extreme technical depth, clear engineering ownership, and quantified impact. Generic resumes that just list keywords without context are instantly rejected.

Your goal is to produce a clean, ATS-optimized resume in plain text based on the candidate's raw resume and the target job description. The final resume MUST be strictly under 2 pages.

Your output must have EXACTLY the following four sections in this order:

PROFESSIONAL SUMMARY
SKILLS
EXPERIENCE
EDUCATION

---

SECTION-BY-SECTION INSTRUCTIONS:

PROFESSIONAL SUMMARY
- Write a highly compelling, results-driven 2–4 sentence summary.
- Start with a strong professional hook that incorporates the target job title (e.g., "Results-driven Senior AI Engineer with 6+ years of experience...") and years of experience.
- Directly highlight 2–3 advanced technical competencies or architectural achievements that align with the core requirements of the Job Description.
- Draw ONLY on facts present in the resume. Do NOT invent experience.
- DO NOT use any emojis, checkmarks, icons, or smileys.

SKILLS
- List technical and professional skills present in the candidate's raw resume.
- Group them logically into categories (e.g., Languages, Frameworks & Libraries, Tools & Platforms, Methodologies).
- Within each category, prioritize the tools that appear in the job description or keyword targets by listing them first.
- Format skills as a simple plain text categorized list, e.g.:
    Languages: Python, Java, Go
    Frameworks: PyTorch, FastAPI, Django
- DO NOT use bullet points (•) or lists in the SKILLS section.
- Normalize terms to match the Job Description's keywords (e.g., use "PostgreSQL" instead of "Postgres").
- Do NOT add skills the candidate does not have.

EXPERIENCE
- Include EVERY job from the candidate's original resume. Do NOT drop any role.
- Each job entry MUST have exactly one header line formatted EXACTLY as:
    Company Name | Job Title | Start Month Year – End Month Year (or Present)
  followed by 3–5 bullet points starting with •
- DO NOT repeat company headers. List all bullet points and projects for a single job entry under ONE single header line. Do NOT split a single job entry into multiple headers.
- For each bullet point, follow the XYZ formula: "Accomplished [X] as measured by [Y], by doing [Z]" using strong active verbs and specific engineering detail.
- DO NOT use weak words like "assisted", "helped", "responsible for", "participated in". Start every bullet with ownership verbs: "Spearheaded", "Architected", "Engineered", "Optimized", "Scaled", "Pioneered".
- Every experience bullet MUST contain a quantitative metric showing impact (e.g., latency reduction ms/%, compute cost reduction %, throughput scaling, or data processing volumes).
- Highlight specific architectural details (e.g., distributed model training using PyTorch FSDP/DeepSpeed, serving optimization with vLLM/Triton, LoRA/QLoRA fine-tuning hyperparameters, or chunking/embedding RAG strategies).
- Rewrite accomplishments to match the core duties listed in the Job Description.
- DO NOT invent, change, or fake the actual company names, schools, or employment dates/tenures.
- DO NOT use any emojis, checkmarks, icons, or smileys.

EDUCATION
- List degrees in reverse chronological order.
- Format: Degree, Field of Study | Institution | Year
- Keep it brief — 1–2 lines per degree.
- Include certifications or relevant coursework on a separate line under the degree if present.
- DO NOT use any emojis, checkmarks, icons, or smileys.

---

FORMATTING RULES (CRITICAL — Output is parsed by code):
- Section headers MUST be exactly: PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, EDUCATION (ALL-CAPS, no other text on that line)
- Job entry headers MUST use the pipe separator: Company | Title | Date
- Bullet points MUST start with • (bullet character)
- Ensure a single blank line between sections and between job entries
- Output ONLY the plain resume text. Do NOT wrap in markdown fences (no ```), and do NOT include any preamble, explanations, or introductory/concluding notes.
- The first line of output must be the candidate's Full Name (if found in the resume).
- The second line must contain contact info on one line (e.g., email | phone | LinkedIn | location).
"""

REFINE_PROMPT = """You are a precise resume line editor.

You are given a resume and a REFINEMENT INSTRUCTION that targets ONE specific line or bullet.
Apply the instruction to that line only and return the COMPLETE, UNCHANGED resume with just that line updated.

Strict rules:
- Change ONLY the specific line described in the instruction.
- DO NOT remove any other line, bullet, section, job, or project.
- DO NOT reorder anything.
- Preserve the exact formatting: ALL-CAPS headers, • bullets, | separators, blank lines.
- Output ONLY the full resume text. No explanations, no markdown fences.
"""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def build_tailored_resume_docx(
    resume_text: str,
    job_description: str,
    candidate_name: str = "",
    job_title: str = "",
    company: str = "",
    matched_keywords: list[str] = None,
    missing_keywords: list[str] = None
) -> tuple[bytes, str]:
    """
    Structures + keyword-aligns the resume via LLM, then renders as DOCX.
    Returns (docx_bytes, filename).
    """
    if job_description:
        tailored_text = await _llm_rewrite(resume_text, job_description, matched_keywords, missing_keywords)
    else:
        tailored_text = resume_text
    return generate_docx_bytes(tailored_text, candidate_name, job_title, company)


async def _llm_rewrite(
    resume_text: str,
    job_description: str,
    matched_keywords: list[str] = None,
    missing_keywords: list[str] = None
) -> str:
    """
    Single-pass LLM call that both restructures the resume into the canonical
    4-section format AND aligns keywords to the job description.
    """
    resume_truncated = resume_text[:8000] if len(resume_text) > 8000 else resume_text
    jd_truncated = job_description[:8000] if len(job_description) > 8000 else job_description

    user_content = (
        f"CANDIDATE'S RESUME:\n{resume_truncated}\n\n"
        f"JOB DESCRIPTION:\n{jd_truncated}\n\n"
    )
    
    if matched_keywords or missing_keywords:
        user_content += "TARGET KEYWORDS TO NATURALLY INTEGRATE (ONLY IF CANDIDATE HAS RELEVANT EXPERIENCE/SKILLS):\n"
        if matched_keywords:
            user_content += f"- Already matched keywords: {', '.join(matched_keywords)}\n"
        if missing_keywords:
            user_content += f"- Missing keywords to incorporate: {', '.join(missing_keywords)}\n"
        user_content += "\n"

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": STRUCTURE_PROMPT},
            {
                "role": "user",
                "content": user_content,
            },
        ],
        "stream": False,
    }

    async with httpx.AsyncClient(timeout=180.0) as client:
        response = await client.post(f"{OLLAMA_URL}/api/chat", json=payload)
        response.raise_for_status()

    content = response.json().get("message", {}).get("content", "").strip()
    # Strip accidental markdown fences
    content = re.sub(r"^```(?:markdown|text)?|```$", "", content, flags=re.MULTILINE).strip()
    
    # Clean double/redundant bullets
    cleaned_lines = []
    for line in content.splitlines():
        cleaned_line = re.sub(r"^\s*[•\-\*·\s]+[•\-\*·]\s*", "• ", line)
        cleaned_lines.append(cleaned_line)
    return "\n".join(cleaned_lines)


async def _llm_refine(current_resume: str, instruction: str) -> str:
    """Apply a targeted refinement instruction to an already-tailored resume."""
    resume_truncated = current_resume[:6000] if len(current_resume) > 6000 else current_resume

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": REFINE_PROMPT},
            {
                "role": "user",
                "content": (
                    f"CURRENT RESUME:\n{resume_truncated}\n\n"
                    f"REFINEMENT INSTRUCTION:\n{instruction}"
                ),
            },
        ],
        "stream": False,
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(f"{OLLAMA_URL}/api/chat", json=payload)
        response.raise_for_status()

    content = response.json().get("message", {}).get("content", "").strip()
    content = re.sub(r"^```(?:markdown|text)?|```$", "", content, flags=re.MULTILINE).strip()
    return content


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

_SECTION_HEADERS = {
    "PROFESSIONAL SUMMARY", "SUMMARY",
    "WORK EXPERIENCE", "EXPERIENCE", "EMPLOYMENT",
    "PROJECTS", "PROJECT EXPERIENCE",
    "EDUCATION",
    "SKILLS", "TECHNICAL SKILLS",
    "CERTIFICATIONS", "AWARDS", "PUBLICATIONS", "VOLUNTEERING",
}

_COLOR_NAME     = RGBColor(0x1a, 0x1a, 0x2e)   # deep navy  — candidate name
_COLOR_SUBHEAD  = RGBColor(0x16, 0x21, 0x3e)   # dark blue  — section headers
_COLOR_COMPANY  = RGBColor(0x1a, 0x1a, 0x2e)   # near-black — company/role lines
_COLOR_META     = RGBColor(0x77, 0x77, 0x88)   # mid-grey   — dates, contact info
_COLOR_BODY     = RGBColor(0x40, 0x40, 0x40)   # dark grey  — body / bullets
_COLOR_RULE     = RGBColor(0xdd, 0xdd, 0xdd)   # light rule


def _safe(text: str) -> str:
    """Normalise common Unicode punctuation for DOCX compatibility."""
    return (
        text
        .replace("\u2019", "'")
        .replace("\u2018", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2013", "–")
        .replace("\u2014", "—")
        .strip()
    )


def _add_rule(doc: Document, thin: bool = False) -> None:
    """Add a thin horizontal rule paragraph using paragraph border."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3) if not thin else Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4" if thin else "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx_bytes(
    tailored_text: str,
    candidate_name: str = "",
    job_title: str = "",
    company: str = "",
) -> tuple[bytes, str]:
    """
    Renders the structured plain-text resume as a polished DOCX.
    Returns (docx_bytes, smart_filename).
    """
    doc = Document()

    # Page margins — tight for 1-page fit
    for section in doc.sections:
        section.top_margin    = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin   = Inches(0.80)
        section.right_margin  = Inches(0.80)

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    lines = tailored_text.splitlines()
    first_content_line = True  # used to detect name / contact header lines

    for raw_line in lines:
        line = raw_line.rstrip()

        # Blank line → small spacer
        if not line.strip():
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            continue

        stripped = _safe(line.strip())
        upper = stripped.upper()

        # ------------------------------------------------------------------
        # Section header (ALL-CAPS known keyword)
        # ------------------------------------------------------------------
        if upper in _SECTION_HEADERS or any(upper.startswith(h) for h in _SECTION_HEADERS):
            first_content_line = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(stripped.upper())
            run.font.size      = Pt(9.5)
            run.font.bold      = True
            run.font.color.rgb = _COLOR_SUBHEAD
            _add_rule(doc, thin=True)
            continue

        # ------------------------------------------------------------------
        # Candidate name (first non-empty line if no section header seen yet)
        # ------------------------------------------------------------------
        if first_content_line and not "|" in stripped and not stripped.startswith("•"):
            # Check if this looks like a name (no @ symbol, not too long)
            if "@" not in stripped and len(stripped.split()) <= 5 and len(stripped) < 60:
                first_content_line = False
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(stripped)
                run.font.size      = Pt(18)
                run.font.bold      = True
                run.font.color.rgb = _COLOR_NAME
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

        # ------------------------------------------------------------------
        # Contact info line (email / LinkedIn / city — contains @ or |)
        # ------------------------------------------------------------------
        if first_content_line or (
            ("@" in stripped or "linkedin" in stripped.lower() or
             ("|" in stripped and not any(
                 kw in stripped.upper()
                 for kw in {"PRESENT", "JAN","FEB","MAR","APR","MAY","JUN",
                             "JUL","AUG","SEP","OCT","NOV","DEC"}) and
             len(stripped.split("|")) <= 3 and
             len(stripped) < 100))
        ):
            if "@" in stripped or "linkedin" in stripped.lower():
                first_content_line = False
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(stripped)
                run.font.size      = Pt(9)
                run.font.color.rgb = _COLOR_META
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

        first_content_line = False

        # ------------------------------------------------------------------
        # Bullet point
        # ------------------------------------------------------------------
        if stripped.startswith(("•", "-", "*", "·")):
            bullet_text = stripped.lstrip("•-*· ").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.20)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(bullet_text)
            run.font.size      = Pt(10)
            run.font.color.rgb = _COLOR_BODY
            continue

        # ------------------------------------------------------------------
        # Company | Title | Date  (job entry header)
        # ------------------------------------------------------------------
        if "|" in stripped and len(stripped.split("|")) >= 2:
            parts = [s.strip() for s in stripped.split("|")]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after  = Pt(1)
            # Company name bold
            r_company = p.add_run(parts[0])
            r_company.font.bold      = True
            r_company.font.size      = Pt(10.5)
            r_company.font.color.rgb = _COLOR_COMPANY
            # Role + dates in lighter weight
            if len(parts) > 1:
                r_rest = p.add_run("  ·  " + "  ·  ".join(parts[1:]))
                r_rest.font.size      = Pt(10)
                r_rest.font.color.rgb = _COLOR_META
            continue

        # ------------------------------------------------------------------
        # Project / sub-heading line (contains — or –)
        # ------------------------------------------------------------------
        if "—" in stripped or "–" in stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(stripped)
            run.font.bold      = True
            run.font.size      = Pt(10.5)
            run.font.color.rgb = _COLOR_COMPANY
            continue

        # ------------------------------------------------------------------
        # Default body text (summary paragraphs, skill lines, etc.)
        # ------------------------------------------------------------------
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(stripped)
        run.font.size      = Pt(10)
        run.font.color.rgb = _COLOR_BODY

    # ------------------------------------------------------------------
    # Filename: {Company}_{JobTitle}_{YYYYMMDD}.docx
    # ------------------------------------------------------------------
    date_str      = datetime.now().strftime("%Y%m%d")
    company_clean = re.sub(r"[^\w]", "_", company or "Export").strip("_")
    job_clean     = re.sub(r"[^\w]", "_", job_title or "Resume").strip("_")
    filename      = f"{company_clean}_{job_clean}_{date_str}.docx"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), filename


# ---------------------------------------------------------------------------
# AI Optimizer Agent
# ---------------------------------------------------------------------------

AGENT_OPTIMIZER_PROMPT = """You are a Senior Technical Recruiter at a top-tier technology company. You only select resumes that showcase extreme technical depth, clear engineering ownership, and quantified impact.

You are given a current resume (structured in the canonical format), a job description, and a target set of missing keywords/skills that must be integrated.

Your task is to reframe and enhance the projects and bullet points under the candidate's existing WORK EXPERIENCE section (and update the SKILLS list) to naturally weave in these missing skills and showcase matching project work or new engineering ideas relevant to the Job Description.

Target skills/keywords to integrate: {missing_keywords}

Rules:
1. ONLY reframe or expand projects and bullet points under existing jobs.
2. DO NOT change, invent, or fake any company names, universities, or work periods/dates.
3. Every bullet point MUST start with a strong active verb and follow the XYZ formula: "Accomplished [X] as measured by [Y], by doing [Z]" using quantified metrics (latency %, cost %, throughput, utilization) and deep technical details (e.g., distributed PyTorch FSDP, serving with vLLM/Triton, quantization, RAG vector search).
4. DO NOT repeat company headers. List all bullet points and projects for a single job entry under ONE single header line. Do NOT split a single job entry into multiple headers.
5. Format skills as a simple plain text categorized list, e.g.:
    Languages: Python, Java, Go
    Frameworks: PyTorch, FastAPI, Django
   DO NOT use bullet points (•) or lists in the SKILLS section.
6. The entire resume MUST be strictly under 2 pages. Limit bullets per job to 3–5 points.
7. DO NOT use any emojis, checkmarks, icons, or smileys anywhere in the resume text.
8. Format all outputs exactly as the input resume:
   - Headers: PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, EDUCATION
   - Bullet points starting with •
   - Company headers starting with Company Name | Job Title | Dates
9. Output ONLY the resume text. No explanations, introductory notes, or markdown fences.
"""

async def optimize_resume_agent(
    resume_text: str,
    job_description: str,
    target_title: str = "",
    max_passes: int = 9,
    status_callback=None
) -> tuple[str, int, list[dict]]:
    """
    Runs an iterative AI optimization agent to tailor projects and bullets,
    evaluating the industry score after each pass until hitting >= 95% score.
    """
    logs = []
    
    # Pass 1: Initial rewrite/structuring
    msg1 = "Pass 1: Initial structuring & alignment..."
    if status_callback:
        await status_callback(msg1)
    print(f"[Optimizer Agent] {msg1}")
         
    # Exclude matched/missing in first pass to do structure and general alignment
    current_resume = await _llm_rewrite(resume_text, job_description)
    
    # Calculate initial score
    eval_result = calculate_industry_score(current_resume, job_description, target_title)
    current_score = eval_result["score"]
    logs.append({"pass": 1, "score": current_score, "breakdown": eval_result["breakdown"]})
    
    if current_score >= 95:
        msg = f"Target score achieved: {current_score}%!"
        if status_callback:
            await status_callback(msg)
        print(f"[Optimizer Agent] {msg}")
        return current_resume, current_score, logs
        
    for p in range(2, max_passes + 2):
        missing = eval_result["missing"]
        targets = missing[:12]
        
        status_msg = f"Pass {p}: Enhancing project depth and semantic alignment..."
        if targets:
            status_msg = f"Pass {p}: Reframing projects & integrating missing skills ({', '.join(targets[:5])})..."
            
        if status_callback:
            await status_callback(status_msg)
        print(f"[Optimizer Agent] {status_msg}")
            
        target_instr = ", ".join(targets) if targets else "Enhance project technical details, distributed scaling, and ML infrastructure metrics matching the JD."
            
        # Call agent optimizer prompt
        payload = {
            "model": OLLAMA_MODEL,
            "messages": [
                {"role": "system", "content": AGENT_OPTIMIZER_PROMPT.format(missing_keywords=target_instr)},
                {
                    "role": "user",
                    "content": (
                        f"CURRENT RESUME:\n{current_resume}\n\n"
                        f"TARGET JOB DESCRIPTION:\n{job_description}"
                    )
                }
            ],
            "stream": False
        }
        
        async with httpx.AsyncClient(timeout=180.0) as client:
            response = await client.post(f"{OLLAMA_URL}/api/chat", json=payload)
            response.raise_for_status()
            
        content = response.json().get("message", {}).get("content", "").strip()
        content = re.sub(r"^```(?:markdown|text)?|```$", "", content, flags=re.MULTILINE).strip()
        
        if len(content) > 100:
            cleaned_lines = []
            for line in content.splitlines():
                cleaned_line = re.sub(r"^\s*[•\-\*·\s]+[•\-\*·]\s*", "• ", line)
                cleaned_lines.append(cleaned_line)
            current_resume = "\n".join(cleaned_lines)
            
        # Re-evaluate
        eval_result = calculate_industry_score(current_resume, job_description, target_title)
        current_score = eval_result["score"]
        logs.append({"pass": p, "score": current_score, "breakdown": eval_result["breakdown"]})
        
        if current_score >= 95:
            msg = f"Target score achieved in Pass {p}: {current_score}%!"
            if status_callback:
                await status_callback(msg)
            print(f"[Optimizer Agent] {msg}")
            break
            
    return current_resume, current_score, logs


async def optimize_pass_agent(
    resume_text: str,
    job_description: str,
    target_title: str = "",
    pass_num: int = 1,
    missing_keywords: list[str] = None
) -> dict:
    """
    Executes a single, targeted pass of the AI optimizer agent and returns
    the resume text, score, and breakdown.
    """
    if pass_num == 1:
        # Pass 1: Initial rewrite/structuring
        current_resume = await _llm_rewrite(resume_text, job_description)
    else:
        # Pass 2+: Reframing experience projects & integrating missing skills
        targets = missing_keywords[:12] if missing_keywords else []
        target_instr = ", ".join(targets) if targets else "Enhance project technical details, distributed scaling, and ML infrastructure metrics matching the JD."
        
        payload = {
            "model": OLLAMA_MODEL,
            "messages": [
                {"role": "system", "content": AGENT_OPTIMIZER_PROMPT.format(missing_keywords=target_instr)},
                {
                    "role": "user",
                    "content": (
                        f"CURRENT RESUME:\n{resume_text}\n\n"
                        f"TARGET JOB DESCRIPTION:\n{job_description}"
                    )
                }
            ],
            "stream": False
        }
        
        async with httpx.AsyncClient(timeout=180.0) as client:
            response = await client.post(f"{OLLAMA_URL}/api/chat", json=payload)
            response.raise_for_status()
            
        content = response.json().get("message", {}).get("content", "").strip()
        content = re.sub(r"^```(?:markdown|text)?|```$", "", content, flags=re.MULTILINE).strip()
        
        if len(content) > 100:
            current_resume = content
        else:
            current_resume = resume_text
            
    # Clean double/redundant bullets
    cleaned_lines = []
    for line in current_resume.splitlines():
        cleaned_line = re.sub(r"^\s*[•\-\*·\s]+[•\-\*·]\s*", "• ", line)
        cleaned_lines.append(cleaned_line)
    final_resume = "\n".join(cleaned_lines)
    
    # Calculate score
    eval_result = calculate_industry_score(final_resume, job_description, target_title)
    
    return {
        "resume_text": final_resume,
        "score": eval_result["score"],
        "matched": eval_result["matched"],
        "missing": eval_result["missing"],
        "breakdown": eval_result["breakdown"]
    }
